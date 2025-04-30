import os
from pathlib import Path
from docx import Document
from datetime import datetime
from config.settings import WORD_TEMPLATE, ACCEPTABLE_MODIFIERS, ACCEPTABLE_POS

def process_line_items(line_items):
    """Process line items for document placeholders using adapted record format"""
    mapping = {}
    for i, line in enumerate(line_items[:6], start=1):
        modifier_raw = line.get("modifier")
        if isinstance(modifier_raw, list):
            modifier = ",".join([m for m in modifier_raw if m in ACCEPTABLE_MODIFIERS])
        elif isinstance(modifier_raw, str):
            modifier = modifier_raw if modifier_raw in ACCEPTABLE_MODIFIERS else ""
        else:
            modifier = ""
        pos = line.get("pos", "11")
        units = line.get("units", 1)
        try:
            charge = float(str(line.get("charge", 0)).replace("$", "").replace(",", ""))
        except Exception:
            charge = 0.0
        try:
            rate = float(str(line.get("validated_rate", 0)).replace("$", "").replace(",", ""))
        except Exception:
            rate = 0.0
        mapping.update({
            f"<dos{i}>": line.get("date_of_service", ""),
            f"<cpt{i}>": line.get("cpt", "N/A"),
            f"<charge{i}>": "${:,.2f}".format(charge),
            f"<units{i}>": units,
            f"<modifier{i}>": modifier,
            f"<pos{i}>": pos,
            f"<rate{i}>": f"{rate:.2f}",
            f"<alwd{i}>": "${:,.2f}".format(rate),
            f"<paid{i}>": "${:,.2f}".format(rate),
            f"<code{i}>": "85, 125"
        })
    for i in range(len(line_items) + 1, 7):
        mapping.update({
            f"<dos{i}>": "", f"<cpt{i}>": "", f"<charge{i}>": "", f"<units{i}>": "",
            f"<modifier{i}>": "", f"<pos{i}>": "", f"<alwd{i}>": "", f"<paid{i}>": "", f"<code{i}>": ""
        })
    return mapping

def populate_placeholders(doc, mapping):
    """Replace placeholders in a Word document with values"""
    sanitized_mapping = {k: str(v) if v is not None else "" for k, v in mapping.items()}
    # Process paragraphs
    for paragraph in doc.paragraphs:
        for placeholder, value in sanitized_mapping.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)
    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for placeholder, value in sanitized_mapping.items():
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, value)

def generate_document(record, eobr_data, output_folders):
    """Generate Word document for an EOBR record using adapted record format"""
    data = record.get("data", {})
    patient_info = data.get("patient_info", {})
    provider_info = data.get("provider_info", {})
    billing_address = provider_info.get("Billing_Address", {})
    line_items = data.get("line_items", [])

    # Patient info
    patient_name = patient_info.get("PatientName", "N/A")
    dob = patient_info.get("Patient_DOB", "")
    injury_date = patient_info.get("Patient_Injury_Date", "")

    # Provider info
    billing_name = provider_info.get("Billing_Name", "N/A")
    billing_address1 = billing_address.get("Address", "N/A")
    billing_city = billing_address.get("City", "N/A")
    billing_state = billing_address.get("State", "N/A")
    billing_zip = billing_address.get("Postal_Code", "N/A")
    tin = provider_info.get("TIN", "N/A")
    npi = provider_info.get("NPI", "N/A")

    # Provider ref
    provider_ref = record.get("billing_info", {}).get("patient_account_no", "N/A")

    # Order number logic
    order_id = record.get("order_id", "")
    filemaker_record_number = patient_info.get("FileMaker_Record_Number", "")
    if order_id and str(order_id).startswith("ORD"):
        order_no = order_id
    else:
        order_no = filemaker_record_number or order_id or "N/A"

    # Total paid: sum validated_rate from line_items
    total_paid = 0.0
    for item in line_items:
        try:
            total_paid += float(str(item.get("validated_rate", 0)).replace("$", "").replace(",", ""))
        except Exception:
            continue

    mapping = {
        "<process_date>": datetime.now().strftime("%Y-%m-%d"),
        "<PatientName>": patient_name,
        "<dob>": dob,
        "<doi>": injury_date,
        "<provider_ref>": provider_ref,
        "<order_no>": order_no,
        "<billing_name>": billing_name,
        "<billing_address1>": billing_address1,
        "<billing_address2>": "",  # Not present in new format
        "<billing_city>": billing_city,
        "<billing_state>": billing_state,
        "<billing_zip>": billing_zip,
        "<TIN>": tin,
        "<NPI>": npi,
        "<total_paid>": "${:,.2f}".format(total_paid),
    }
    mapping.update(process_line_items(line_items))
    doc = Document(WORD_TEMPLATE)
    populate_placeholders(doc, mapping)
    eobr_file_name = f"EOBR_{eobr_data['EOBR Number']}"
    docx_output = os.path.join(output_folders['docs'], f"{eobr_file_name}.docx")
    doc.save(docx_output)
    return docx_output, None  # Return None for pdf_path since we're not generating PDFs